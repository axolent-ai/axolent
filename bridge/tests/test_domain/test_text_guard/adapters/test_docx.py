"""Tests for the DOCX adapter (via TextGuardService).

These tests are skipped if python-docx is not installed.
"""

from __future__ import annotations

from pathlib import Path

import pytest

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from application.text_guard_service import TextGuardService

pytestmark = pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")


@pytest.fixture
def tg_service() -> TextGuardService:
    """Text Guard service instance."""
    return TextGuardService()


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample DOCX with ASCII umlauts."""
    doc = Document()
    doc.add_paragraph("Das ist fuer dich.")
    doc.add_paragraph("Natuerlich ist das moeglich.")
    doc.add_paragraph("This is correct English.")
    filepath = tmp_path / "test.docx"
    doc.save(str(filepath))
    return filepath


@pytest.fixture
def clean_docx(tmp_path: Path) -> Path:
    """Create a DOCX with correct umlauts."""
    doc = Document()
    doc.add_paragraph("Das ist für dich.")
    filepath = tmp_path / "clean.docx"
    doc.save(str(filepath))
    return filepath


class TestDocxCheck:
    """Tests for checking DOCX files."""

    def test_finds_issues(
        self, tg_service: TextGuardService, sample_docx: Path
    ) -> None:
        """Detects ASCII umlauts in DOCX paragraphs."""
        issues = tg_service.check_file(sample_docx, "de")
        assert len(issues) >= 2

    def test_clean_file(self, tg_service: TextGuardService, clean_docx: Path) -> None:
        """Clean DOCX returns no issues."""
        issues = tg_service.check_file(clean_docx, "de")
        assert issues == []


class TestDocxFix:
    """Tests for fixing DOCX files."""

    def test_fixes_in_place(
        self, tg_service: TextGuardService, sample_docx: Path
    ) -> None:
        """Fixes ASCII umlauts in DOCX in-place."""
        changed = tg_service.fix_file(sample_docx, "de")
        assert changed is True

        # Verify the fix
        doc = Document(str(sample_docx))
        texts = [p.text for p in doc.paragraphs]
        assert "für" in texts[0]
        assert "Natürlich" in texts[1]

    def test_no_change_on_clean(
        self, tg_service: TextGuardService, clean_docx: Path
    ) -> None:
        """Clean DOCX is not modified."""
        changed = tg_service.fix_file(clean_docx, "de")
        assert changed is False
